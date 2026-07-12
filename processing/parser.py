"""
文件解析 — 把上传的 PDF / Word / Markdown / txt 解析为纯文本

PDF  → PyMuPDF (import fitz)
docx → python-docx
md/txt → 直接按 UTF-8 读取（Markdown 本身即近似纯文本，检索无需渲染）
"""
from pathlib import Path

from utils.logger import log


SUPPORTED_TYPES = {"pdf", "docx", "md", "txt"}
# 兼容的扩展名 → 归一化类型
_EXT_MAP = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".md": "md",
    ".markdown": "md",
    ".txt": "txt",
    ".text": "txt",
}


def detect_file_type(filename: str) -> str:
    """从文件名后缀识别类型，未知返回空串"""
    return _EXT_MAP.get(Path(filename).suffix.lower(), "")


def _parse_pdf(path: Path) -> str:
    import fitz  # PyMuPDF
    parts = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def _parse_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # 顺带抽取表格文本（企业文档常用表格）
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_text(path: Path) -> str:
    # 尽量宽容地读取，避免个别乱码字节导致整篇失败
    return path.read_text(encoding="utf-8", errors="ignore")


def parse_file(file_path: Path, file_type: str = "") -> str:
    """
    解析文件为纯文本。
    Args:
        file_path: 磁盘路径
        file_type: pdf/docx/md/txt，留空则按扩展名推断
    Returns:
        解析出的纯文本（已 strip）
    Raises:
        ValueError: 不支持的类型
        Exception:  解析库抛出的底层异常（由调用方捕获并标记 failed）
    """
    file_path = Path(file_path)
    ftype = (file_type or detect_file_type(file_path.name)).lower()
    if ftype not in SUPPORTED_TYPES:
        raise ValueError(f"不支持的文件类型: {ftype or file_path.suffix}（支持 {sorted(SUPPORTED_TYPES)}）")

    if ftype == "pdf":
        text = _parse_pdf(file_path)
    elif ftype == "docx":
        text = _parse_docx(file_path)
    else:  # md / txt
        text = _parse_text(file_path)

    text = (text or "").strip()
    log.info(f"[parser] 解析完成 {file_path.name} ({ftype}): {len(text)} 字符")
    return text
